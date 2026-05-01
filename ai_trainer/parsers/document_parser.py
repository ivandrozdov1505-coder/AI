"""
Парсер документов: PDF, DOCX, DOC, XLSX
"""

from pathlib import Path
from typing import Dict, Any, List, Optional
from .base_parser import BaseParser
import logging

logger = logging.getLogger(__name__)


class DocumentParser(BaseParser):
    """Парсер для документов"""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__(config)
        self.supported_extensions = ['.pdf', '.docx', '.doc', '.xlsx']
    
    def can_parse(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.supported_extensions
    
    def parse(self, file_path: str) -> List[Dict[str, Any]]:
        """Парсинг документа"""
        if not self.validate_file(file_path):
            logger.error(f"Файл не найден: {file_path}")
            return []
        
        suffix = Path(file_path).suffix.lower()
        
        if suffix == '.pdf':
            return self._parse_pdf(file_path)
        elif suffix == '.docx':
            return self._parse_docx(file_path)
        elif suffix == '.doc':
            return self._parse_doc(file_path)
        elif suffix == '.xlsx':
            return self._parse_xlsx(file_path)
        else:
            logger.warning(f"Неподдерживаемый формат: {suffix}")
            return []
    
    def _parse_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """Парсинг PDF с использованием PyMuPDF"""
        data = []
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            
            for page_num, page in enumerate(doc):
                text = page.get_text()
                if text.strip():
                    data.append({
                        'text': text,
                        'source': str(file_path),
                        'page': page_num + 1,
                        'type': 'pdf'
                    })
            
            doc.close()
            logger.debug(f"PDF: загружено {len(data)} страниц")
            
        except ImportError:
            logger.error("PyMuPDF не установлен. Установите: pip install PyMuPDF")
        except Exception as e:
            logger.error(f"Ошибка парсинга PDF: {e}")
        
        return data
    
    def _parse_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """Парсинг DOCX"""
        data = []
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Параграфы
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    data.append({
                        'text': para.text.strip(),
                        'source': str(file_path),
                        'paragraph': i + 1,
                        'type': 'docx'
                    })
            
            # Таблицы
            for t_idx, table in enumerate(doc.tables):
                for r_idx, row in enumerate(table.rows):
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        data.append({
                            'text': ' | '.join(row_data),
                            'source': str(file_path),
                            'table': t_idx + 1,
                            'row': r_idx + 1,
                            'type': 'docx_table'
                        })
            
            logger.debug(f"DOCX: загружено {len(data)} элементов")
            
        except ImportError:
            logger.error("python-docx не установлен. Установите: pip install python-docx")
        except Exception as e:
            logger.error(f"Ошибка парсинга DOCX: {e}")
        
        return data
    
    def _parse_doc(self, file_path: str) -> List[Dict[str, Any]]:
        """Парсинг DOC (устаревший формат)"""
        logger.warning("Формат .doc устарел. Конвертируйте в .docx для лучшей поддержки.")
        # Попытка прочитать как текст (может работать для некоторых файлов)
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            lines = [line.strip() for line in content.split('\n') if line.strip()]
            return [{
                'text': line,
                'source': str(file_path),
                'type': 'doc_fallback'
            } for line in lines]
        except Exception as e:
            logger.error(f"Ошибка парсинга DOC: {e}")
            return []
    
    def _parse_xlsx(self, file_path: str) -> List[Dict[str, Any]]:
        """Парсинг XLSX"""
        data = []
        try:
            from openpyxl import load_workbook
            
            wb = load_workbook(file_path, read_only=True)
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                
                for row_idx, row in enumerate(sheet.iter_rows(values_only=True)):
                    # Пропускаем пустые строки
                    if all(cell is None for cell in row):
                        continue
                    
                    row_data = {
                        f'col_{i}': cell for i, cell in enumerate(row) if cell is not None
                    }
                    
                    if row_data:
                        row_data['source'] = str(file_path)
                        row_data['sheet'] = sheet_name
                        row_data['row'] = row_idx + 1
                        row_data['type'] = 'xlsx'
                        data.append(row_data)
            
            wb.close()
            logger.debug(f"XLSX: загружено {len(data)} строк")
            
        except ImportError:
            logger.error("openpyxl не установлен. Установите: pip install openpyxl")
        except Exception as e:
            logger.error(f"Ошибка парсинга XLSX: {e}")
        
        return data
